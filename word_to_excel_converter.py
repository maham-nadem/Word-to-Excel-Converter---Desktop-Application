#!/usr/bin/env python3
"""
Word to Excel Converter - Desktop Application
Converts Word documents to Excel spreadsheets with a custom logo and drag-and-drop interface
"""

import sys
import os
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QLabel, QPushButton, QFileDialog, 
                             QMessageBox, QProgressBar, QTextEdit)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QMimeData
from PyQt6.QtGui import QDragEnterEvent, QDropEvent, QPixmap, QIcon
import pandas as pd
from docx import Document
import traceback

class ConversionWorker(QThread):
    """Worker thread for conversion to keep UI responsive"""
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    finished = pyqtSignal(bool, str)
    
    def __init__(self, input_file, output_file):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        
    def run(self):
        try:
            self.status.emit("Loading Word document...")
            self.progress.emit(20)
            
            # Load Word document
            doc = Document(self.input_file)
            
            self.status.emit("Extracting tables...")
            self.progress.emit(40)
            
            # Extract tables from Word
            all_tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:  # Only add non-empty tables
                    df = pd.DataFrame(table_data)
                    all_tables.append((f"Table_{table_idx + 1}", df))
            
            self.status.emit("Creating Excel file...")
            self.progress.emit(70)
            
            # Create Excel file with multiple sheets
            with pd.ExcelWriter(self.output_file, engine='openpyxl') as writer:
                if all_tables:
                    for sheet_name, df in all_tables:
                        df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                else:
                    # If no tables found, create a simple sheet with document text
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    if paragraphs:
                        df = pd.DataFrame(paragraphs, columns=["Content"])
                        df.to_excel(writer, sheet_name="Document Content", index=False)
            
            self.progress.emit(100)
            self.status.emit("Conversion complete!")
            self.finished.emit(True, f"Successfully converted to:\n{self.output_file}")
            
        except Exception as e:
            self.finished.emit(False, f"Error: {str(e)}\n\n{traceback.format_exc()}")

class WordToExcelConverter(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()
        
    def initUI(self):
        self.setWindowTitle("Word to Excel Converter")
        self.setMinimumSize(600, 500)
        
        # Set application icon (if you have one)
        if os.path.exists("icon.ico"):
            self.setWindowIcon(QIcon("icon.ico"))
        elif os.path.exists("icon.png"):
            self.setWindowIcon(QIcon("icon.png"))
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # Logo area
        logo_layout = QHBoxLayout()
        logo_label = QLabel()
        
        # Try to load logo from various possible locations
        logo_paths = ["logo.png", "logo.ico", "icon.png", "icon.ico", "app_logo.png"]
        logo_loaded = False
        
        for path in logo_paths:
            if os.path.exists(path):
                pixmap = QPixmap(path)
                if not pixmap.isNull():
                    scaled_pixmap = pixmap.scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio, 
                                                  Qt.TransformationMode.SmoothTransformation)
                    logo_label.setPixmap(scaled_pixmap)
                    logo_loaded = True
                    break
        
        if not logo_loaded:
            # Create a text logo if no image found
            logo_label.setText("📄➡️📊")
            logo_label.setStyleSheet("font-size: 48px; color: #2c3e50;")
        
        logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        logo_layout.addWidget(logo_label)
        
        # Title
        title_label = QLabel("Word to Excel Converter")
        title_label.setStyleSheet("font-size: 24px; font-weight: bold; color: #3498db; margin: 10px;")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        main_layout.addLayout(logo_layout)
        main_layout.addWidget(title_label)
        
        # Drag and drop area
        self.drop_area = QLabel()
        self.drop_area.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.drop_area.setMinimumHeight(150)
        self.drop_area.setStyleSheet("""
            QLabel {
                border: 3px dashed #3498db;
                border-radius: 10px;
                background-color: #ecf0f1;
                font-size: 14px;
                color: #7f8c8d;
                margin: 10px;
            }
            QLabel:hover {
                border-color: #2980b9;
                background-color: #d5dbdb;
            }
        """)
        self.drop_area.setText("Drag and Drop Word File Here\n\n-or-\n\nClick to Select File")
        
        # Enable drag and drop
        self.drop_area.setAcceptDrops(True)
        self.drop_area.dragEnterEvent = self.dragEnterEvent
        self.drop_area.dropEvent = self.dropEvent
        self.drop_area.mousePressEvent = self.select_file
        
        main_layout.addWidget(self.drop_area)
        
        # Selected file display
        self.file_label = QLabel("No file selected")
        self.file_label.setStyleSheet("color: #34495e; padding: 5px;")
        self.file_label.setWordWrap(True)
        main_layout.addWidget(self.file_label)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                                        
                border: 1px solid #bdc3c7;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #3498db;
                border-radius: 5px;
            }
        """)
        main_layout.addWidget(self.progress_bar)
        
        # Status text
        self.status_text = QTextEdit()
        self.status_text.setVisible(False)
        self.status_text.setMaximumHeight(100)
        self.status_text.setStyleSheet("background-color: #f8f9fa; color: #2c3e50;")
        main_layout.addWidget(self.status_text)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.convert_btn = QPushButton("Convert to Excel")
        self.convert_btn.setEnabled(False)
        self.convert_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                font-size: 14px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
            }
        """)
        self.convert_btn.clicked.connect(self.convert_file)
        
        self.clear_btn = QPushButton("Clear")
        self.clear_btn.setEnabled(False)
        self.clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                font-size: 14px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
            }
        """)
        self.clear_btn.clicked.connect(self.clear_selection)
        
        button_layout.addWidget(self.convert_btn)
        button_layout.addWidget(self.clear_btn)
        main_layout.addLayout(button_layout)
        
        # Footer with info
        footer_label = QLabel("Supports .docx files | Tables are preserved in Excel sheets")
        footer_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        footer_label.setStyleSheet("color: #95a5a6; margin-top: 10px;")
        main_layout.addWidget(footer_label)
        
        # Store selected file
        self.selected_file = None
        self.conversion_thread = None
        
    def dragEnterEvent(self, event: QDragEnterEvent):
        """Handle drag enter events"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.drop_area.setStyleSheet("""
                QLabel {
                    border: 3px dashed #27ae60;
                    border-radius: 10px;
                    background-color: #e8f8f5;
                    font-size: 14px;
                    color: #7f8c8d;
                    margin: 10px;
                }
            """)
    
    def dropEvent(self, event: QDropEvent):
        """Handle drop events"""
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        if files:
            self.process_file(files[0])
        
        # Reset drop area style
        self.drop_area.setStyleSheet("""
            QLabel {
                border: 3px dashed #3498db;
                border-radius: 10px;
                background-color: #ecf0f1;
                font-size: 14px;
                color: #7f8c8d;
                margin: 10px;
            }
        """)
    
    def select_file(self, event):
        """Handle click on drop area to select file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Word File", "", "Word Documents (*.docx);;All Files (*)"
        )
        if file_path:
            self.process_file(file_path)
    
    def process_file(self, file_path):
        """Process selected file"""
        if file_path.lower().endswith('.docx'):
            self.selected_file = file_path
            self.file_label.setText(f"Selected: {os.path.basename(file_path)}")
            self.convert_btn.setEnabled(True)
            self.clear_btn.setEnabled(True)
        else:
            QMessageBox.warning(self, "Invalid File", "Please select a valid .docx Word document.")
    
    def clear_selection(self):
        """Clear current selection"""
        self.selected_file = None
        self.file_label.setText("No file selected")
        self.convert_btn.setEnabled(False)
        self.clear_btn.setEnabled(False)
        self.progress_bar.setVisible(False)
        self.status_text.setVisible(False)
        self.progress_bar.setValue(0)
    
    def convert_file(self):
        """Convert Word to Excel"""
        if not self.selected_file:
            return
        
        # Ask for output location
        output_file, _ = QFileDialog.getSaveFileName(
            self, "Save Excel File", "", "Excel Files (*.xlsx)"
        )
        
        if not output_file:
            return
        
        # Ensure .xlsx extension
        if not output_file.endswith('.xlsx'):
            output_file += '.xlsx'
        
        # Show progress UI
        self.progress_bar.setVisible(True)
        self.status_text.setVisible(True)
        self.status_text.clear()
        self.convert_btn.setEnabled(False)
        self.clear_btn.setEnabled(False)
        
        # Start conversion in separate thread
        self.conversion_thread = ConversionWorker(self.selected_file, output_file)
        self.conversion_thread.progress.connect(self.progress_bar.setValue)
        self.conversion_thread.status.connect(self.update_status)
        self.conversion_thread.finished.connect(self.conversion_finished)
        self.conversion_thread.start()
    
    def update_status(self, message):
        """Update status text"""
        self.status_text.append(message)
    
    def conversion_finished(self, success, message):
        """Handle conversion completion"""
        self.convert_btn.setEnabled(True)
        self.clear_btn.setEnabled(True)
        
        if success:
            QMessageBox.information(self, "Success", message)
        else:
            QMessageBox.critical(self, "Error", message)
            self.status_text.append(message)

def main():
    app = QApplication(sys.argv)
    app.setApplicationName("Word to Excel Converter")
    app.setOrganizationName("YourName")
    
    # Set application icon if exists
    if os.path.exists("icon.ico"):
        app.setWindowIcon(QIcon("icon.ico"))
    
    window = WordToExcelConverter()
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
#!/usr/bin/env python3
"""
Installation script for Word to Excel Converter
Creates desktop shortcut and prepares the application
"""

import os
import sys
import subprocess
import platform
from pathlib import Path

def create_desktop_shortcut_pyshortcuts():
    """Create desktop shortcut using pyshortcuts library"""
    try:
        from pyshortcuts import Shortcut
        
        # Get current directory
        current_dir = os.path.dirname(os.path.abspath(__file__))
        script_path = os.path.join(current_dir, "word_to_excel_converter.py")
        
        # Find icon
        icon_path = None
        for icon_name in ["logo.ico", "icon.ico", "logo.png"]:
            potential_icon = os.path.join(current_dir, icon_name)
            if os.path.exists(potential_icon):
                icon_path = potential_icon
                break
        
        # Create shortcut
        sc = Shortcut(
            script=script_path,
            name="Word to Excel Converter",
            description="Convert Word documents to Excel spreadsheets",
            icon=icon_path if icon_path else "",
            terminal=False
        )
        
        # Save to desktop
        sc.save(desktop=True)
        print("✓ Desktop shortcut created using pyshortcuts")
        return True
        
    except ImportError:
        print("pyshortcuts not installed, trying alternative method...")
        return False
    except Exception as e:
        print(f"Error with pyshortcuts: {e}")
        return False

def create_desktop_shortcut_winshell():
    """Create desktop shortcut using winshell (Windows only)"""
    if platform.system() != 'Windows':
        return False
    
    try:
        import winshell
        from win32com.client import Dispatch
        
        current_dir = os.path.dirname(os.path.abspath(__file__))
        script_path = os.path.join(current_dir, "word_to_excel_converter.py")
        python_exe = sys.executable
        
        # Create shortcut target (python script.py)
        target = f'"{python_exe}" "{script_path}"'
        
        # Find icon
        icon_path = ""
        for icon_name in ["logo.ico", "icon.ico"]:
            potential_icon = os.path.join(current_dir, icon_name)
            if os.path.exists(potential_icon):
                icon_path = potential_icon
                break
        
        desktop = winshell.desktop()
        shortcut_path = os.path.join(desktop, "Word to Excel Converter.lnk")
        
        shell = Dispatch('WScript.Shell')
        shortcut = shell.CreateShortCut(shortcut_path)
        shortcut.Targetpath = python_exe
        shortcut.Arguments = f'"{script_path}"'
        shortcut.WorkingDirectory = current_dir
        if icon_path:
            shortcut.IconLocation = icon_path
        shortcut.save()
        
        print("✓ Desktop shortcut created using winshell")
        return True
        
    except ImportError:
        return False
    except Exception as e:
        print(f"Error with winshell: {e}")
        return False

def create_batch_file():
    """Create a batch file to run the application"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    script_path = os.path.join(current_dir, "word_to_excel_converter.py")
    batch_path = os.path.join(current_dir, "WordToExcelConverter.bat")
    
    batch_content = f"""@echo off
echo Starting Word to Excel Converter...
cd /d "{current_dir}"
python "{script_path}"
pause
"""
    
    with open(batch_path, "w") as f:
        f.write(batch_content)
    
    print("✓ Batch file created")
    return batch_path

def create_manual_instructions():
    """Create manual instructions if automatic shortcut creation fails"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    script_path = os.path.join(current_dir, "word_to_excel_converter.py")
    
    instructions = f"""
===============================================
MANUAL SHORTCUT CREATION INSTRUCTIONS
===============================================

To create a desktop shortcut manually:

1. Right-click on your desktop
2. Select "New" → "Shortcut"
3. For the location, enter:
   {sys.executable} "{script_path}"
4. Click "Next"
5. Name the shortcut: Word to Excel Converter
6. Click "Finish"

To add your custom icon:
1. Right-click the new shortcut
2. Select "Properties"
3. Click "Change Icon"
4. Browse to your icon file (logo.ico)
5. Click "OK" twice

Your application is located at:
{script_path}

Python interpreter: {sys.executable}
Working directory: {current_dir}
"""
    
    instructions_path = os.path.join(current_dir, "SHORTCUT_INSTRUCTIONS.txt")
    with open(instructions_path, "w") as f:
        f.write(instructions)
    
    print("⚠ Could not create shortcut automatically")
    print(f"✓ Manual instructions saved to: {instructions_path}")
    return instructions_path

def main():
    print("\n" + "="*50)
    print("Word to Excel Converter - Installation")
    print("="*50 + "\n")
    
    # Check Python version
    print(f"Python version: {sys.version}")
    if sys.version_info < (3, 6):
        print("⚠ Warning: Python 3.6 or higher is recommended")
    
    # Install required packages
    print("\n📦 Installing required packages...")
    packages = ["PyQt6", "pandas", "python-docx", "openpyxl", "pyshortcuts"]
    
    for package in packages:
        try:
            __import__(package.replace("-", "_").split(">=")[0])
            print(f"  ✓ {package} already installed")
        except ImportError:
            print(f"  Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    
    # Try to create desktop shortcut
    print("\n🖥️ Creating desktop shortcut...")
    
    shortcut_created = False
    
    # Try pyshortcuts first
    if create_desktop_shortcut_pyshortcuts():
        shortcut_created = True
    
    # Try winshell if pyshortcuts failed (Windows only)
    if not shortcut_created and platform.system() == 'Windows':
        if create_desktop_shortcut_winshell():
            shortcut_created = True
    
    # Create batch file as backup
    batch_path = create_batch_file()
    
    # If all automatic methods failed, provide manual instructions
    if not shortcut_created:
        create_manual_instructions()
    
    print("\n" + "="*50)
    print("✅ Installation Complete!")
    print("="*50)
    
    print("\nWhat to do next:")
    print("  1. Double-click the desktop shortcut to run the application")
    print("  OR")
    print(f"  2. Run the batch file: {batch_path}")
    
    if not shortcut_created:
        print("\n📝 Follow the instructions in SHORTCUT_INSTRUCTIONS.txt")
    
    print("\n📁 Your logo files should be in this folder:")
    print(f"   {os.path.dirname(os.path.abspath(__file__))}")
    print("\n   The application will look for: logo.png, logo.ico, icon.png, or icon.ico")
    
    input("\nPress Enter to exit...")

if __name__ == "__main__":
    main()    